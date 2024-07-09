from docx import Document
import io


def write_message(msg):
    lines = msg.split("\n")
    doc = Document()
    doc.add_heading(lines[0], level=1)
    doc.add_paragraph("\n".join(lines[1:]))
    doc.save("docx_file.docx")


def create_word_document(content):
    doc = Document()
    lines = content.split("\n")
    doc.add_heading(lines[0])
    doc.add_paragraph(lines[1:])
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
